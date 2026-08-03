"""
Genera el CV en formato Harvard (ES + EN) del candidato configurado en .env.
Uso: .venv/bin/python3 profile/generate_cv.py
Produce: profile/cv_es.docx  y  profile/cv_en.docx
PDF: libreoffice --headless --convert-to pdf --outdir profile profile/cv_es.docx profile/cv_en.docx
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv
import copy

TEXT_WIDTH_INCHES = 7.66  # 8.5" page - 0.42" * 2 margins

load_dotenv(Path(__file__).resolve().parent.parent / ".env")


def _env_join(*keys: str) -> str:
    """Une los valores presentes de esas vars con el separador del encabezado.

    Los enlaces van sin esquema en el CV impreso ("github.com/x", no
    "https://github.com/x"), asi que se recorta al vuelo.
    """
    vals = (
        os.getenv(k, "").strip().removeprefix("https://").removeprefix("http://").removeprefix("www.").rstrip("/")
        for k in keys
    )
    return "  •  ".join(v for v in vals if v)


# El encabezado del CV lleva el nombre legal completo; CANDIDATE_NAME es el que
# se usa en formularios y cartas, y puede ser la version corta.
NAME = (os.getenv("CV_FULL_NAME") or os.getenv("CANDIDATE_NAME", "")).strip()
LOCATION = os.getenv("CANDIDATE_LOCATION", "").strip() or ", ".join(
    v for v in (os.getenv("CANDIDATE_CITY", "").strip(), os.getenv("CANDIDATE_COUNTRY", "").strip()) if v
)
# El CV sale a empresas reales: mejor reventar aca que emitir un PDF sin nombre.
assert NAME, "Falta CANDIDATE_NAME en el .env"

CONTACT_LINE_1 = "  •  ".join(v for v in (LOCATION, os.getenv("GMAIL_USER", "").strip(), os.getenv("CANDIDATE_PHONE", "").strip()) if v)
CONTACT_LINE_2 = _env_join("CANDIDATE_LINKEDIN", "CANDIDATE_GITHUB", "CANDIDATE_WEBSITE")


# ── helpers ────────────────────────────────────────────────────────────────────

def _add_right_tab(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(TEXT_WIDTH_INCHES * 1440)))
    tabs.append(tab)
    pPr.append(tabs)


def _add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_spacing(paragraph, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(12)


def _set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.42)
        section.right_margin = Inches(0.42)


def _base_run(paragraph, text, bold=False, italic=False, size=11):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


# ── building blocks ─────────────────────────────────────────────────────────────

def add_name(doc, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=0, after=2)
    _base_run(p, name, bold=True, size=16)


def add_contact(doc, line1, line2=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=0, after=6)
    _base_run(p, line1, size=9.5)
    if line2:
        p.add_run("\n")
        _base_run(p, line2, size=9.5)


def add_section(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bottom_border(p)
    _set_spacing(p, before=8, after=2)
    _base_run(p, title.upper(), bold=True, size=11)


def add_org_line(doc, org, location, space_before=4):
    """Bold org name [TAB] location (right-aligned)."""
    p = doc.add_paragraph()
    _add_right_tab(p)
    _set_spacing(p, before=space_before, after=0)
    _base_run(p, org, bold=True)
    _base_run(p, "\t")
    _base_run(p, location)


def add_role_line(doc, role, date):
    """Italic role [TAB] date (right-aligned)."""
    p = doc.add_paragraph()
    _add_right_tab(p)
    _set_spacing(p, before=0, after=0)
    _base_run(p, role, italic=True)
    _base_run(p, "\t")
    _base_run(p, date)


def add_body(doc, text, space_before=0):
    p = doc.add_paragraph()
    _set_spacing(p, before=space_before, after=0)
    _base_run(p, text)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _set_spacing(p, before=0, after=0)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.2)
    pf.first_line_indent = Inches(-0.15)
    _base_run(p, text)


def add_skills_line(doc, label, value):
    p = doc.add_paragraph()
    _set_spacing(p, before=1, after=0)
    _base_run(p, label + ": ", bold=True)
    _base_run(p, value)


# ── ES document ────────────────────────────────────────────────────────────────

def build_es():
    doc = Document()
    _set_margins(doc)

    # remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    add_name(doc, NAME)
    add_contact(doc, CONTACT_LINE_1, CONTACT_LINE_2)

    # ── PERFIL ──
    add_section(doc, "Perfil")
    add_body(
        doc,
        "Desarrollador fullstack con sistemas en producción para clientes reales en Ecuador "
        "(gestión clínica, facturación, inventario multisucursal). "
        "Especializado en Django/DRF y React/Next.js, con pruebas automatizadas y CI/CD, "
        "y formación complementaria en ciberseguridad.",
    )

    # ── EXPERIENCIA ──
    # Experiencia va ANTES de Educación: el rol de cofundador y los sistemas en
    # producción pesan más que la etiqueta de "9no semestre".
    add_section(doc, "Experiencia")

    add_org_line(doc, "Facilito", "Guayaquil, Ecuador")
    add_role_line(doc, "Co-fundador & Desarrollador Fullstack", "Feb. 2026 – presente")
    add_bullet(doc, "Desarrollé backend con Django 5 + DRF + PostgreSQL para plataforma SaaS multitenancy del sector gastronómico.")
    add_bullet(doc, "Integré asistente IA con Claude (Anthropic) para consultas de inventario, gestión de pedidos y reportes automáticos.")
    add_bullet(doc, "Construí app móvil con Flutter + Riverpod que consume la API REST con dio y almacenamiento seguro de sesión.")
    add_bullet(doc, "Diseñé multitenancy por queryset: manager for_tenant() sobre modelo base abstracto que fuerza el aislamiento de datos antes de llegar al ViewSet.")

    add_org_line(doc, "Freelance", "Guayaquil, Ecuador", space_before=6)
    add_role_line(doc, "Desarrollador Fullstack", "Jul. 2025 – presente")
    add_bullet(doc, "Desarrollé sistema de gestión clínica Tia Glenda (React + MUI + Flask + PostgreSQL, ~130 componentes) en producción, con RBAC de 3 roles y 23 suites de tests.")
    add_bullet(doc, "Construí plataforma de pedidos con FastAPI + SQLAlchemy + Alembic, en producción en novamicktools.com.")
    add_bullet(doc, "Implementé la integración con el SRI de Ecuador para un e-commerce: generación del XML, clave de acceso de 49 dígitos y envío de sobres SOAP a recepción y autorización (ambiente de pruebas).")
    add_bullet(doc, "Implementé SimuladorPreguntas universitario (React + Express 5 + PostgreSQL) con roles diferenciados, medidas anti-trampa y despliegue en Docker multi-stage con usuario no-root y nginx.")
    add_bullet(doc, "Desarrollé Taller App: gestión para taller mecánico (Express + Sequelize + PDFKit + React), con 5 suites de tests Jest + Supertest.")

    add_org_line(doc, "Área de Nivelación — Universidad de Guayaquil", "Guayaquil, Ecuador", space_before=6)
    add_role_line(doc, "Practicante de Desarrollo de Software", "Feb. 2026 – Jul. 2026")
    add_bullet(doc, "Lideré el desarrollo frontend, backend, pruebas y deploy del SimuladorPreguntas para uso institucional.")
    add_bullet(doc, "Automaticé procesos internos de registro y generación de reportes mediante scripts Python.")

    # ── EDUCACIÓN ──
    add_section(doc, "Educación")

    add_org_line(doc, "Universidad de Guayaquil", "Guayaquil, Ecuador", space_before=4)
    add_role_line(doc, "Ingeniería en Software — 9no semestre (de 10)", "2021 – presente")
    add_body(doc, "Promedio académico: 9.01 / 10  (escala 0–10, aprobación desde 7)")

    add_org_line(doc, "Google / Coursera", "En línea", space_before=4)
    add_role_line(doc, "Certificado Profesional de Ciberseguridad", "2026 – en progreso")

    add_org_line(doc, "Google / Coursera", "En línea", space_before=4)
    add_role_line(doc, "Certificado Profesional de UX Design", "2026 – en progreso")

    add_org_line(doc, "Academia de Ciberseguridad Hacker Mentor", "En línea", space_before=4)
    add_role_line(doc, "Curso de Ethical Hacking — Red Team (8 horas)", "2023")

    # ── PROYECTOS ──
    add_section(doc, "Proyectos")

    add_org_line(doc, "RestoVentas", "github.com/Mickaell22/restoventas-backend")
    add_role_line(doc, "App de ventas para restaurante con pedidos por voz", "NestJS · TypeORM · React Native · IA")
    add_bullet(doc, "Construí backend NestJS 11 con TypeORM y migraciones, autenticación Passport-JWT con guards y DTOs validados con class-validator.")
    add_bullet(doc, "Desarrollé app React Native + Expo con Zustand para el estado global y captura de audio para tomar pedidos hablados.")
    add_bullet(doc, "Implementé parseo de pedidos en lenguaje natural con LLM + STT, saneando la salida del modelo contra el catálogo real de productos para no confiar en ella a ciegas.")

    add_org_line(doc, "MotoVox", "github.com/Mickaell22/MotoVox", space_before=6)
    add_role_line(doc, "App de comunicación por voz para motociclistas", "Flutter · C nativo · FFI · Sockets")
    add_bullet(doc, "Integré código C nativo con Flutter vía FFI (compilado con NDK para ARM64) para filtrado de ruido en tiempo real sin overhead de JVM.")
    add_bullet(doc, "Implementé audio peer-to-peer sobre WiFi local con sockets TCP crudos (tcpNoDelay para minimizar latencia) y descubrimiento de salas por broadcast UDP, sin servidor intermediario.")

    add_org_line(doc, "QR Shield", "github.com/Mickaell22/qr-shield", space_before=6)
    add_role_line(doc, "Motor de detección de QR maliciosos (tesis)", "Python · FastAPI · Tests")
    add_bullet(doc, "Construí API REST en FastAPI con una capa de heurísticas de URL para detectar quishing, cubierta con tests automatizados.")
    add_bullet(doc, "Diseñé la arquitectura por capas de análisis L1–L5, con las integraciones de reputación de dominios planificadas como siguiente etapa.")

    add_org_line(doc, "ApplyJob", "github.com/Mickaell22/ApplyJob", space_before=6)
    add_role_line(doc, "Pipeline de agregación y matching de ofertas con LLMs", "Python · Playwright · LLMs")
    add_bullet(doc, "Construí agregación de 13 fuentes heterogéneas (APIs JSON, RSS y HTML) normalizadas a un esquema común, con Playwright headless para las que requieren render JS.")
    add_bullet(doc, "Implementé el matching contra el perfil y la generación de texto con LLMs; reordenar el prompt dejó el 75% constante y cacheable como prefijo, y un test de regresión cubre la extracción del stack.")

    # ── HABILIDADES ──
    add_section(doc, "Habilidades")

    add_skills_line(doc, "Backend", "Python, Django/DRF, FastAPI, Flask, Node.js, Express, NestJS, C#, Java, PHP")
    add_skills_line(doc, "Frontend", "React, Next.js, TypeScript, Tailwind CSS, Vite, Material UI, Zustand, Redux Toolkit")
    add_skills_line(doc, "Móvil", "Flutter, Dart, Riverpod, React Native, Expo, Kotlin")
    add_skills_line(doc, "Bases de datos", "PostgreSQL, SQLAlchemy/Alembic, TypeORM, Prisma, Sequelize, Firebase, SQLite")
    add_skills_line(doc, "Testing", "pytest, Jest, Vitest, Testing Library, Supertest, JUnit + Mockito")
    add_skills_line(doc, "DevOps", "Git, Docker (multi-stage), CI/CD con GitHub Actions, Linux, VPS (Railway), nginx")
    add_skills_line(doc, "Ciberseguridad", "Kali Linux, Nmap, Wireshark, fundamentos de ethical hacking")
    add_skills_line(doc, "Idiomas", "Español (nativo), Inglés (B2 — lectura, escritura técnica y conversacional)")

    return doc


# ── EN document ────────────────────────────────────────────────────────────────

def build_en():
    doc = Document()
    _set_margins(doc)

    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    add_name(doc, NAME)
    add_contact(doc, CONTACT_LINE_1, CONTACT_LINE_2)

    # ── PROFILE ──
    add_section(doc, "Profile")
    add_body(
        doc,
        "Fullstack developer with systems running in production for real clients in Ecuador "
        "(clinical management, order platform, multi-branch inventory). "
        "Focused on Django/DRF and React/Next.js, with automated testing and CI/CD, "
        "and complementary training in cybersecurity.",
    )

    # ── EXPERIENCE ──
    add_section(doc, "Experience")

    add_org_line(doc, "Facilito", "Guayaquil, Ecuador")
    add_role_line(doc, "Co-founder & Fullstack Developer", "Feb 2026 – present")
    add_bullet(doc, "Built backend with Django 5 + Django REST Framework + PostgreSQL for a multitenancy SaaS targeting the restaurant industry.")
    add_bullet(doc, "Integrated an AI assistant powered by Claude (Anthropic) for inventory queries, order management, and automated reports.")
    add_bullet(doc, "Developed mobile app with Flutter + Riverpod consuming the REST API with dio, with secure session storage.")
    add_bullet(doc, "Designed queryset-level multitenancy: a for_tenant() manager on an abstract base model that enforces data isolation before the ViewSet is reached.")

    add_org_line(doc, "Freelance", "Guayaquil, Ecuador", space_before=6)
    add_role_line(doc, "Fullstack Developer", "Jul 2025 – present")
    add_bullet(doc, "Built Tia Glenda clinical management system (React + MUI + Flask + PostgreSQL, ~130 components) in production, with 3-role RBAC and 23 test suites.")
    add_bullet(doc, "Built an order management platform with FastAPI + SQLAlchemy + Alembic, running in production at novamicktools.com.")
    add_bullet(doc, "Implemented the integration with Ecuador's SRI tax authority for an e-commerce: XML generation, 49-digit access key, and SOAP envelopes sent to the reception and authorization endpoints (test environment).")
    add_bullet(doc, "Implemented university exam simulator (React + Express 5 + PostgreSQL) with role-based access, anti-cheating measures, and a multi-stage Docker deployment running as a non-root user behind nginx.")
    add_bullet(doc, "Delivered Taller App: workshop management system (Express + Sequelize + PDFKit + React) with 5 Jest + Supertest suites.")

    add_org_line(doc, "Leveling Area — University of Guayaquil", "Guayaquil, Ecuador", space_before=6)
    add_role_line(doc, "Software Development Intern", "Feb 2026 – Jul 2026")
    add_bullet(doc, "Led frontend, backend, testing, and deployment of the SimuladorPreguntas platform for institutional use.")
    add_bullet(doc, "Automated internal registration and report-generation workflows using Python scripts.")

    # ── EDUCATION ──
    add_section(doc, "Education")

    add_org_line(doc, "University of Guayaquil", "Guayaquil, Ecuador", space_before=4)
    add_role_line(doc, "B.Sc. Software Engineering — 9th semester (of 10)", "2021 – present")
    add_body(doc, "GPA: 9.01 / 10  (grading system: 0–10, minimum passing grade: 7)")

    add_org_line(doc, "Google / Coursera", "Online", space_before=4)
    add_role_line(doc, "Professional Certificate in Cybersecurity", "2026 – in progress")

    add_org_line(doc, "Google / Coursera", "Online", space_before=4)
    add_role_line(doc, "Professional Certificate in UX Design", "2026 – in progress")

    add_org_line(doc, "Hacker Mentor Cybersecurity Academy", "Online", space_before=4)
    add_role_line(doc, "Ethical Hacking — Red Team course (8 hours)", "2023")

    # ── PROJECTS ──
    add_section(doc, "Projects")

    add_org_line(doc, "RestoVentas", "github.com/Mickaell22/restoventas-backend")
    add_role_line(doc, "Restaurant sales app with voice-driven ordering", "NestJS · TypeORM · React Native · AI")
    add_bullet(doc, "Built a NestJS 11 backend with TypeORM and migrations, Passport-JWT authentication with guards, and DTOs validated with class-validator.")
    add_bullet(doc, "Developed a React Native + Expo app using Zustand for global state, with audio capture for spoken orders.")
    add_bullet(doc, "Implemented natural-language order parsing with an LLM plus speech-to-text, sanitizing the model output against the real product catalog rather than trusting it blindly.")

    add_org_line(doc, "MotoVox", "github.com/Mickaell22/MotoVox", space_before=6)
    add_role_line(doc, "Voice communication app for motorcycle riders", "Flutter · Native C · FFI · Sockets")
    add_bullet(doc, "Integrated native C code with Flutter via FFI (compiled with the NDK for ARM64) for real-time noise filtering with no JVM overhead.")
    add_bullet(doc, "Implemented peer-to-peer audio over local WiFi using raw TCP sockets (tcpNoDelay to minimize latency) and UDP broadcast for room discovery, with no intermediary server.")

    add_org_line(doc, "QR Shield", "github.com/Mickaell22/qr-shield", space_before=6)
    add_role_line(doc, "Malicious QR code detection engine (thesis)", "Python · FastAPI · Tests")
    add_bullet(doc, "Built a REST API in FastAPI with a URL heuristics layer to detect quishing, covered by automated tests.")
    add_bullet(doc, "Designed the L1-L5 layered analysis architecture, with domain reputation integrations planned as the next stage.")

    add_org_line(doc, "ApplyJob", "github.com/Mickaell22/ApplyJob", space_before=6)
    add_role_line(doc, "Job posting aggregation and matching pipeline with LLMs", "Python · Playwright · LLMs")
    add_bullet(doc, "Built aggregation of 13 heterogeneous sources (JSON APIs, RSS, and HTML) normalized into a common schema, with headless Playwright for JS-rendered ones.")
    add_bullet(doc, "Implemented profile matching and LLM text generation; reordering the prompt left 75% of it constant and cacheable as a prefix, with a regression test covering stack extraction.")

    # ── SKILLS ──
    add_section(doc, "Skills")

    add_skills_line(doc, "Backend", "Python, Django/DRF, FastAPI, Flask, Node.js, Express, NestJS, C#, Java, PHP")
    add_skills_line(doc, "Frontend", "React, Next.js, TypeScript, Tailwind CSS, Vite, Material UI, Zustand, Redux Toolkit")
    add_skills_line(doc, "Mobile", "Flutter, Dart, Riverpod, React Native, Expo, Kotlin")
    add_skills_line(doc, "Databases", "PostgreSQL, SQLAlchemy/Alembic, TypeORM, Prisma, Sequelize, Firebase, SQLite")
    add_skills_line(doc, "Testing", "pytest, Jest, Vitest, Testing Library, Supertest, JUnit + Mockito")
    add_skills_line(doc, "DevOps", "Git, Docker (multi-stage), CI/CD with GitHub Actions, Linux, VPS (Railway), nginx")
    add_skills_line(doc, "Cybersecurity", "Kali Linux, Nmap, Wireshark, ethical hacking fundamentals")
    add_skills_line(doc, "Languages", "Spanish (native), English (B2 — reading, technical writing, conversational)")

    return doc


# ── main ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import pathlib

    out = pathlib.Path(__file__).parent

    # nombres que consumen CV_PATH / CV_PATH_EN del .env
    es_path = out / "cv_es.docx"
    en_path = out / "cv_en.docx"

    build_es().save(es_path)
    print(f"✓ {es_path}")

    build_en().save(en_path)
    print(f"✓ {en_path}")
