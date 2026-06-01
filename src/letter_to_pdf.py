"""
Convierte cartas .txt a PDF via python-docx + LibreOffice.
Uso: python src/letter_to_pdf.py output/cartas/07_Canonical_SWE_Python_Cloud.txt
"""

import subprocess
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def txt_to_pdf(txt_path: str) -> str:
    txt = Path(txt_path)
    out_dir = txt.parent
    docx_path = out_dir / (txt.stem + ".docx")
    pdf_path = out_dir / (txt.stem + ".pdf")

    doc = Document()

    # Margenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    lines = txt.read_text(encoding="utf-8").splitlines()

    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line if line else "")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if not line:
            p.paragraph_format.space_before = Pt(6)

    doc.save(str(docx_path))

    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        check=True,
        capture_output=True,
    )

    docx_path.unlink()
    print(f"PDF generado: {pdf_path}")
    return str(pdf_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python src/letter_to_pdf.py <ruta_carta.txt>")
        sys.exit(1)
    txt_to_pdf(sys.argv[1])
